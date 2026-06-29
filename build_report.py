"""Build the HW1 report (DOCX) with question labels, formulas, explanations,
and the figures/outputs embedded."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text='', bold=False, italic=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def formula(text):
    # Display a formula centred, in a serif-ish font for readability.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p

def mono(text):
    # Monospaced block for pasted code outputs.
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(8)
    return p

def figure(path, width=5.8, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)
            c.paragraph_format.space_after = Pt(10)
    else:
        para(f'[missing figure: {path}]', italic=True)

D = 'HW1_data'

# ======================================================== Title page
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('50.007 Machine Learning — Homework 1')
r.bold = True; r.font.size = Pt(20)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Report'); r.font.size = Pt(14)
nm = doc.add_paragraph(); nm.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nm.add_run('Name: Somesh Agrawal      Student ID: ____________________')
r.font.size = Pt(11)
para('', space_after=2)

# ======================================================== Q1
heading('Q1 — Perceptron (Linear Classification)', level=1)
para('We implement the perceptron with an offset and train on train_1_5.csv '
     '(columns: symmetry, intensity, label), evaluating on test_1_5.csv. As '
     'required, the data is NOT shuffled — instances are visited sequentially. '
     'A point is misclassified when:')
formula('yᵢ (θ · xᵢ + θ₀) ≤ 0')
para('and on each misclassification we update:')
formula('θ ← θ + yᵢ xᵢ ,    θ₀ ← θ₀ + yᵢ')

heading('Q1(a) — 1 epoch (one pass over the training set)', level=2)
mono('theta  = [-2.4483233, -5.88190688]\n'
     'offset = 0.0\n'
     'test accuracy = 0.9671')
para('After a single pass, θ = [-2.4483, -5.8819], offset = 0, and the test '
     'accuracy is 96.71%.')

heading('Q1(b) — 5 epochs', level=2)
mono('theta  = [-2.05895878, -8.83521606]\n'
     'offset = 1.0\n'
     'test accuracy = 0.9671')
para('After 5 epochs, θ = [-2.0590, -8.8352], offset = 1, and the test accuracy '
     'is still 96.71%. The data is close to linearly separable, so even one pass '
     'classifies ~97% of the test set correctly; extra epochs change the weights '
     'but not the test accuracy here.')

heading('Q1(c) — How to run', level=2)
para('The submitted notebook FirstName_LastName_HW1.ipynb contains the code under '
     'the "Q1" heading. Run from the project folder (which contains HW1_data/) and '
     'execute the cells top to bottom. perceptron(X_train, y_train, epochs) returns '
     '(theta, offset); accuracy(X_test, y_test, theta, offset) returns the test '
     'accuracy. Data is read from HW1_data/perceptron/.')

# ======================================================== Q2
heading('Q2 — Linear and Polynomial Regression', level=1)
para('Inputs hw1x.dat, targets hw1y.dat. Training error is reported as the '
     'empirical risk (mean squared error):')
formula('R(θ) = (1/n) Σᵢ (xᵢᵀθ − yᵢ)²')

heading('Q2(a) — Closed-form linear regression', level=2)
para('We prepend a column of 1s (bias) and solve the normal equations:')
formula('θ = (XᵀX)⁻¹ XᵀY')
mono('theta (bias, slope): [3.2447, 1.7816]\n'
     'training MSE       : 1.1625')
figure(os.path.join(D, 'linear-and-polynomial-regression','linear_regression.png'), width=5.2,
       caption='Q2(a): data with the closed-form linear fit.')

heading('Q2(b) — Gradient descent and stochastic gradient descent', level=2)
para('Both use learning rate η = 0.01 for 5 epochs. Batch GD makes one update per '
     'epoch using the full-data gradient; SGD makes one update per data point '
     '(shuffling each epoch, seeded for reproducibility).')
mono('Batch GD final theta: [0.2413, 0.2703]   min training MSE: 22.3745\n'
     'SGD      final theta: [3.0091, 1.9637]   min training MSE: 1.1763')
para('With only 5 epochs and this small step size, batch GD (5 updates total) '
     'barely moves from zero, so its error is still large. SGD makes 200 updates '
     'per epoch, so it reaches close to the closed-form solution (θ ≈ [3.24, 1.78], '
     'MSE ≈ 1.16) in the same 5 epochs.')
figure(os.path.join(D, 'linear-and-polynomial-regression','gd_convergence.png'), width=5.0,
       caption='Q2(b): training MSE per epoch for GD vs SGD (dotted = closed form).')

heading('Q2(c) — Polynomial regression', level=2)
para('PolyRegress(x, y, d) builds the design matrix [1, x, x², …, xᵈ] and solves '
     'it with the same closed form as part (a).')
para('Quadratic (degree-2) fit:', bold=True)
mono('degree-2 theta: [3.5808, 0.7834, 0.4966]\n'
     'degree-2 MSE  : 1.1406')
figure(os.path.join(D, 'linear-and-polynomial-regression','poly_degree2.png'), width=5.0,
       caption='Q2(c): degree-2 polynomial fit.')
para('Training error for degrees 1–15:', bold=True)
mono('d  1 -> 1.1625    d  6 -> 1.1227    d 11 -> 1.4402\n'
     'd  2 -> 1.1406    d  7 -> 1.1125    d 12 -> 6.4399\n'
     'd  3 -> 1.1402    d  8 -> 1.1076    d 13 -> 9.3212\n'
     'd  4 -> 1.1292    d  9 -> 1.1058    d 14 -> 5.7676\n'
     'd  5 -> 1.1278    d 10 -> 1.1052    d 15 -> 100.5221')
figure(os.path.join(D, 'linear-and-polynomial-regression','poly_error_vs_degree.png'), width=5.0,
       caption='Q2(c): training MSE vs polynomial degree.')
figure(os.path.join(D, 'linear-and-polynomial-regression','poly_all_degrees.png'), width=6.2,
       caption='Q2(c): polynomial fits for degrees 1–15.')
para('After which order does the error get worse? The error decreases up to '
     'degree 10, then increases starting at degree 11 (1.1052 → 1.4402). In theory '
     'training error can never increase with more terms; here it does only for '
     'numerical reasons: the columns x, x², …, x¹⁵ become almost collinear, so XᵀX '
     'is nearly singular (ill-conditioned) and np.linalg.inv loses accuracy. Using '
     'np.linalg.lstsq or np.linalg.pinv would keep the error flat.')

# ======================================================== Q3
heading('Q3 — Ridge Regression', level=1)
para('Inputs hw1_ridge_x.dat (already include a column of 1s), targets '
     'hw1_ridge_y.dat. The first 10 rows are the validation set, the last 40 are '
     'the training set. The closed-form ridge solution is:')
formula('θ̂ = (n λ I + XᵀX)⁻¹ XᵀY')
para('where n = 40 (number of training points).')

heading('Q3(a) — θ for λ = 0.15', level=2)
mono('theta = [-0.57939825,  1.15033028,  0.04934122, -1.59867896]')

heading('Q3(b) — Choosing λ with the validation set', level=2)
para('We sweep λ on a log scale (log₁₀λ from 0 down to −4.9, the grid given in the '
     'question) and plot training vs validation loss. The chosen λ minimises the '
     'validation loss.')
figure(os.path.join(D, 'ridge-regression','ridge_loss_vs_lambda.png'), width=5.2,
       caption='Q3(b): log training loss (red) and log validation loss (blue) vs log₁₀λ.')
mono('lambda minimising validation loss = 0.012589  (≈ 10^-1.9)')
para('The training loss rises monotonically with λ, while the validation loss dips '
     'to a minimum near λ ≈ 0.0126 and then rises again — that minimum is the λ we '
     'pick.')

# ======================================================== Q4
heading('Q4 — K-Means on an Image', level=1)
para('kmeans-image.txt holds one RGB vector per pixel (516×407 image). We run our '
     'own K-Means with K = 8, the given initial centroids, and squared Euclidean '
     'distance. Each iteration assigns every pixel to its nearest centroid and '
     'recomputes each centroid as the mean of its pixels, stopping when the '
     'assignments stop changing.')
para('Convergence:', bold=True)
para('The algorithm converged after 48 iterations. The total within-cluster '
     'squared error drops fast at first — from ≈ 1.29×10⁹ at iteration 1 to '
     '≈ 1.18×10⁸ — and then flattens out (full per-iteration error is printed in '
     'the notebook).')
para('Final centroids and pixel counts:', bold=True)
mono('  k        R        G        B     pixels\n'
     '  0   241.23   238.63   233.86       4930\n'
     '  1   194.41   136.33    90.94      15190\n'
     '  2   136.27    61.09    10.10      52535\n'
     '  3     0.00   255.00     0.00          0\n'
     '  4   157.29    97.59    51.43      22075\n'
     '  5     0.00     0.00   255.00          0\n'
     '  6    78.93    37.11    13.07      40365\n'
     '  7    25.98    23.24    23.61      74917')
figure(os.path.join(D, 'k-means','kmeans_result.png'), width=6.0,
       caption='Q4: original image (left) vs K-Means 8-colour quantised image (right).')
para('Clusters 3 (pure green) and 5 (pure blue) end up with 0 pixels. The image is '
     'a portrait with warm skin tones, a dark suit and a light background, so there '
     'are no pure-green or pure-blue pixels for those centroids to attract — they '
     'stay at their initial positions. The other 6 centroids cover the real colours '
     'of the image.')

out = 'Somesh_Agrawal_HW1_Report.docx'
doc.save(out)
print('wrote', out)
